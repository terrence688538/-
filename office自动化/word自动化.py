from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt       #磅数
from docx.oxml.ns import qn      #中文格式
from docx.shared import Inches  #图片尺寸
import time

#逻辑是先Document文档的属性   然后是paragraph段落的属性      然后是文字的属性   不能混淆

price=input('请输入今日价格')
company_list=['客户1','客户2','客户3','客户4','客户5','客户6','客户7','客户8','客户9','客户10']
today1=time.strftime('%Y-%m-%d',time.localtime())
today2=time.strftime('%Y/%m/%d',time.localtime())
today3=time.strftime('%Y{y}%m{m}%d{d}',time.localtime()).format(y='年',m='月',d='日')
print(today1)

for i in company_list:
    document=Document()
    #styles为样式，Normal是默认  但这个只能设置英文字体
    document.styles['Normal'].font.name=u'宋体'          #后面字符串以 Unicode 格式 进行编码，一般用在中文字符串前面，防止因为源码储存格式问题，导致再次使用时出现乱码。
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')          #不必理解，遇到就把宋体改成想要的字体就好了
    document.add_picture(r'C:\Users\Mechrevo\Desktop\7.png',width=Inches(6))             #增加图片

    p1=document.add_paragraph()
    p1.alignment=WD_ALIGN_PARAGRAPH.CENTER                #对齐方式为居中，没有的话默认左对齐
    run1=p1.add_run('关于下达%s产品价格的通知'%(today3))
    run1.font.name='微软雅黑'                                        #英文
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')      #中文
    run1.font.size=Pt(21)
    run1.font.bold=True                     #加粗
    p1.space_after=Pt(5)                    #段后
    p1.space_before=Pt(5)                   #段前

    p2=document.add_paragraph()
    run2=p2.add_run(i+':')
    run2.font.name='仿宋_GB2312'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size=Pt(16)
    run2.font.bold=True

    p3=document.add_paragraph()
    run3=p3.add_run('  根据公司安排，为提供优质客户服务，我单位拟定了今日黄金价格为%s元，特此通知。'%price)
    run3.font.name = '仿宋_GB2312'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size=Pt(16)
    run3.font.bold=True

    table=document.add_table(rows=3,cols=3,style='Table Grid')
    table.cell(0,0).merge(table.cell(0,2))             #0,0 合并到（0，2）
    table_run1=table.cell(0,0).paragraphs[0].add_run('XX产品报价表')
    table_run1.font.name=u'隶书'
    table_run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书')
    table.cell(0,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    table.cell(1,0).text='日期'
    table.cell(1, 1).text = '价格'
    table.cell(1, 2).text = '备注'
    table.cell(2, 0).text = today3
    table.cell(2, 1).text = str(price)
    table.cell(2, 2).text = ''

    p4 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('（联系人：小杨    电话：18888888888）')
    run4.font.name = '仿宋_GB2312'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)
    run4.font.bold = True

    document.add_page_break()           #在翻一页

    p5=document.add_paragraph()
    run5=p5.add_run('此处是广告')

    document.save(r'C:\Users\Mechrevo\Desktop\6/%s-价格通知.docx' % i)


#读取
document_1=Document(r'C:\Users\Mechrevo\Desktop\作业.docx')                             #读取文字
all_paragraphs_1=document_1.paragraphs
for paragraphs in all_paragraphs_1:
    print(paragraphs.text)

document_2=Document(r'C:\Users\Mechrevo\Desktop\客户8-价格通知.docx')                   #读取表格
all_tables=document_2.tables
for table in all_tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

import zipfile                                                                          #有表格有文字
import re
word=zipfile.ZipFile(r'C:\Users\Mechrevo\Desktop\客户8-价格通知 (1).docx')
xml=word.read('word/document.xml').decode('utf-8')
print(xml)
a=re.findall('<w:t>(.*?)</w:t>',xml)


#封装函数节约时间
document_w=Document()
document_w.styles['Normal'].font.name=u'黑体'
document_w.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

def add_context(context):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(context))
    r.font.size=Pt(16)
    p.space_after=Pt(5)
    p.space_before = Pt(5)

change='哈士奇'
add_context('心塞吹牛不i册为彼此比本次%s恶气丝虫病你吃撒从怒'%change)
add_context('尖下巴元素为本次恶霸v此吹牛五吹牛欸我冰女ui我v不i认为')
add_context('秘书外侧u并不如哦哦亲摩擦去弄i哦v你问包容你我v迥旺vivow0')
add_context('迷哦测侬额我v呢哦v弄从都叫我v弄i额外v你哦那我哦i前面从')
add_context('八年卒七八次完全被此前纽夫妾%s腹腔内粗去玩你想死粗牛'%change)
document_w.save()


#word转pdf
from win32com.client import Dispatch,constants,gencache
doc_path='C:/Users/Mechrevo/Desktop/客户10-价格通知.docx'
pdf_path='C:/Users/Mechrevo/Desktop/客户10-价格通知.pdf'
gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}',0,8,4)
wd=Dispatch("Word.Application")
doc=wd.Documents.Open(doc_path,ReadOnly=1)
doc.ExportAsFixedFormat(pdf_path,constants.wdExportFormatPDF,Item=constants.wdExportDocumentWithMarkup,CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
wd.Quit(constants.wdDoNotSaveChanges)



#读取pdf
from io import StringIO                                                         #数据的无结构传递
from pdfminer.pdfinterp import PDFResourceManager,process_pdf                   #pdf资源管理器、pdf进程
from pdfminer.converter import TextConverter                                    #pdf文字转换
from pdfminer.layout import LAParams                                            #层的参数
# "r" 以读方式打开，只能读文件 ， 如果文件不存在，会发生异常
# "w" 以写方式打开，只能写文件， 如果文件不存在，创建该文件；如果文件已存在，先清空，再打开文件   
# "rb" 以二进制读方式打开，只能读文件 ， 如果文件不存在，会发生异常      
# "wb" 以二进制写方式打开，只能写文件， 如果文件不存在，创建该文件；如果文件已存在，先清空，再打开文件
pdf_file=open('C:/Users/Mechrevo/Desktop/作业12.pdf','rb')               #只需要改动这地方
rsrcmgr=PDFResourceManager()
retstr=StringIO()
laparams=LAParams()
device=TextConverter(rsrcmgr=rsrcmgr,outfp=retstr,laparams=laparams)
process_pdf(rsrcmgr=rsrcmgr,device=device,fp=pdf_file)
device.close()
content=retstr.getvalue()
retstr.close()
pdf_file.close()
print(content)